"""
handlers/docx.py
Extracts text from .docx files using python-docx.
Includes body paragraphs, tables, headers, and footers.
"""

from pathlib import Path
from typing import Any


def extract(filepath: Path) -> dict[str, Any]:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return {
            "text": None,
            "metadata": {},
            "error": "python-docx not installed. Run: pip install python-docx",
        }

    try:
        doc = Document(str(filepath))
    except Exception as exc:
        return {
            "text": None,
            "metadata": {},
            "error": f"Failed to open DOCX: {exc}",
        }

    parts: list[str] = []

    # Core properties / metadata
    metadata: dict[str, Any] = {}
    try:
        cp = doc.core_properties
        for attr in ("author", "title", "subject", "created", "modified"):
            val = getattr(cp, attr, None)
            if val:
                metadata[attr] = str(val)
    except Exception:
        pass

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables — each cell on its own line, rows separated by pipe
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells
            )
            if row_text.strip():
                parts.append(row_text)

    # Headers and footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for para in hf.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    metadata["paragraph_count"] = len(doc.paragraphs)
    metadata["table_count"] = len(doc.tables)

    return {
        "text": "\n\n".join(parts),
        "metadata": metadata,
        "error": None,
    }
